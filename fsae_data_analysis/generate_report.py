import sys
import os
import subprocess

import docx
from docx.shared import Inches


n = len(sys.argv)
if n > 2:
    print("Too many arguments, this script takes 1 filepath (title of folder with graphs).")
    sys.exit()

test_name = sys.argv[1]

# enter the folder if exists
# else run the data_analysis file which generates processes the data and creates the photos
try:
    os.chdir(test_name)
except: 
    subprocess.run(["python3", "data_analysis.py", test_name + ".csv"])
    os.chdir(test_name)

# Create an instance of a word document
doc = docx.Document()
 
# Title
doc.add_heading(test_name.replace('_', ' ') + ' Report', 0)


def insert_img(image, width): # try except pass shortened
    try:
        doc.add_picture(image, width=Inches(width))
    except:
        doc.add_paragraph("No" + image[:-4])
        pass

# Steering
doc.add_heading('Steering', 3)
insert_img('Steering_Wheel_Angle.png', 5)

# Long Slip
doc.add_heading('Longitudinal Slip', 3)
insert_img('Long_Slip.png', 5)

# Acceleration
doc.add_heading('Acceleration', 3)
insert_img('GG.png', 5)
insert_img('Acc.png', 5)

# Wheel Speed
doc.add_heading('Wheel Speed', 3)
insert_img('Wheel_Speed.png', 5)

# Torque
doc.add_heading('Torque', 3)
insert_img('Engine_Torque.png', 5)
insert_img('Torque_Feedback.png', 5)

# Dampers
doc.add_heading('Dampers', 3)
insert_img('Damper_Pos.png', 5)
insert_img('Damper_Vel.png', 5)
insert_img('Damper_Vel_Hist.png', 5)

'''
# Roll
doc.add_picture('Left_Turn_Roll.png')
doc.add_picture('Right_Turn_Roll.png')
doc.add_picture('Roll_Time.png')
'''

# Batteries
doc.add_heading('Batteries', 3)
insert_img('Battery_Temps.png', 5)
insert_img('Battery_Volts.png', 5)

# Coolant
doc.add_heading('Coolant', 3)
insert_img('Coolant_Flow.png', 5)
insert_img('Coolant_Temp.png', 5)
insert_img('Coolant_Pres.png', 5)

# Save doc
doc.save(test_name + '_Report'+ '.docx')